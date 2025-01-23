import os
import telebot
import speech_recognition
from pydub import AudioSegment
import DB as db
from tabulate import tabulate
import openpyxl
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx import Document
from openpyxl import Workbook


# Ниже нужно вставить токен, который дал BotFather при регистрации
# Пример: token = ''
token = '6416692197:AAG5sa9mOys1kwru9ISEvQwtyMuwqRloEVE'  # <<< Ваш токен

bot = telebot.TeleBot(token)


def oga2wav(filename):
    # Конвертация формата файлов
    new_filename = filename.replace('.oga', '.wav')
    audio = AudioSegment.from_file(filename)
    audio.export(new_filename, format='wav')
    return new_filename


def recognize_speech(oga_filename):
    # Перевод голоса в текст + удаление использованных файлов
    wav_filename = oga2wav(oga_filename)
    recognizer = speech_recognition.Recognizer()

    with speech_recognition.WavFile(wav_filename) as source:
        wav_audio = recognizer.record(source)

    text = recognizer.recognize_google(wav_audio, language='ru')

    if os.path.exists(oga_filename):
        os.remove(oga_filename)

    if os.path.exists(wav_filename):
        os.remove(wav_filename)

    return text


def download_file(bot, file_id):
    # Скачивание файла, который прислал пользователь
    file_info = bot.get_file(file_id)
    downloaded_file = bot.download_file(file_info.file_path)
    filename = file_id + file_info.file_path
    filename = filename.replace('/', '_')
    with open(filename, 'wb') as f:
        f.write(downloaded_file)
    return filename

def extract_value_from_table(doc_path, table_index, row_index, col_index):
    # Открываем документ
    doc = Document(doc_path)
    # Получаем таблицу по индексу
    table = doc.tables[table_index]
    # Извлекаем значение из указанной ячейки
    value = table.cell(row_index, col_index).text
    return value


def insert_value_to_table(doc_path, table_index, row_index, col_index, value, new_file_name):
    # Открываем документ
    doc = Document(doc_path)
    # Получаем таблицу по индексу
    table = doc.tables[table_index]
    # Вставляем значение в указанную ячейку
    table.cell(row_index, col_index).text = value
    # Сохраняем изменения
    doc.save(new_file_name)
    print(row_index, ' row index')
    return True



@bot.message_handler(commands=['start'])
def say_hi(message):
    # Функция, отправляющая "Привет" в ответ на команду /start
    bot.send_message(message.chat.id, f'Привет, {message.from_user.first_name}.\nЗапиши голосовое сообщение и получи расшифровку.\n\nЧтобы сформировать окончательный документ введи команду /final')

    # Проверка на наличие юзера в базе
    # работа с БД
    connection = db.create_connection("DEFECTS.sqlite", 'Подключение для проверки на сущ юзера')
    checkUser=f"SELECT COUNT(telegram_id) FROM defects WHERE telegram_id={message.from_user.id} GROUP BY telegram_id"
    ch=db.execute_read_query(connection,checkUser,'Поиск юзера в базе')
    if ch:
        print('Юзер найден')
    else:
        print('нету юзера')
        zapros=f"INSERT INTO defects ('telegram_id') VALUES ({message.chat.id})"
        db.execute_query(connection, zapros, 'создание записи в БД для юзера')
    connection.close()


# defects = dict() #Словарь для сбора замечаний, необходимо перевести в БД
# d = 0

@bot.message_handler(content_types=['voice'])
def transcript(message):
    bot.send_message(message.chat.id, 'Идет расшифровка...')
    # Функция, отправляющая текст в ответ на голосовое
    filename = download_file(bot, message.voice.file_id)
    text = recognize_speech(filename)
    bot.send_message(message.chat.id, text)

    #работа с БД
    connection = db.create_connection("DEFECTS.sqlite", '2')

    take_defect = f"SELECT * FROM defects WHERE telegram_id={message.chat.id}"
    text_defect ,= db.execute_read_query(connection, take_defect, '3')
    # print(text_defect[2])
    text = str(text)
    text_defect = str(text_defect[2])
    text_for_db = text +' $$$ ' + text_defect

    input_defect = f"UPDATE defects SET defect='{text_for_db}' WHERE telegram_id={message.chat.id}"
    db.execute_query(connection, input_defect, 'внесение в БД замечание')

    connection.close()
    #конец работы с БД

    # global d
    # d += 1
    # defects[f'Замечание {d}'] = text

@bot.message_handler(content_types=['document'])
def handle_document(message):
    # Проверяем, что файл - это документ Word
    if message.document.mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
        # Получае файл
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        with open(message.document.file_name,'wb') as new_file:
            new_file.write(downloaded_file)

        bot.send_message(message.chat.id, 'Файл загружен успешно')
        bot.send_message(message.chat.id, message.document.file_name)

        # Параметры для извлечения значения
        source_doc = message.document.file_name # Путь к исходному документу

        table_index_source = 0  # Индекс таблицы в исходном документе
        row_index_source = 0     # Индекс строки в исходной таблице
        col_index_source = 0     # Индекс столбца в исходной таблице


        # Извлекаем значение из исходного документа
        value_to_insert = extract_value_from_table(source_doc, table_index_source, row_index_source, col_index_source)
        
        act_number = value_to_insert.split('№ ', 1) #################### ЭТО ЗНАЧЕНИЕ ДЛЯ ФОРМИРОВАНИЯ ОТЧЕТА ####################
        bot.send_message(message.chat.id, act_number[1])

         # Нас интересует первая таблица в документе
        doc = Document(source_doc)
        table = doc.tables[0]
        
        # Получаем количество строк в таблице
        row_count = len(table.rows)
        # начиная с 3-й строки (2я для кода) и минус 8 строк с конца
        print(f'Количество строк в таблице: {row_count}')
        equipment_count = row_count - (3 + 5)
        print(f'Количество оборудования в таблице: {equipment_count}')

        all_equipments_for_insert = []
        all_kks_numbers_for_insert = [] 
        all_safety_classes_for_insert = []
        all_of_each_counts_for_insert = []
        
        j=0
        for i in range (2, equipment_count):
            equip_name = extract_value_from_table(source_doc, table_index_source, i, 1)
            kks_number = extract_value_from_table(source_doc, table_index_source, i, 2)
            safety_class = extract_value_from_table(source_doc, table_index_source, i, 3)
            each_count = extract_value_from_table(source_doc, table_index_source, i, 4)
            j+=1
            
#################### ЭТО ЗНАЧЕНИЯ ДЛЯ ФОРМИРОВАНИЯ ОТЧЕТА ####################
            all_equipments_for_insert.append(f'{j}. {equip_name}.\n')
            all_kks_numbers_for_insert.append(f'{j}. {kks_number}.\n')
            all_safety_classes_for_insert.append(f'{j}. {safety_class}.\n')
            all_of_each_counts_for_insert.append(f'{j}. {each_count}.\n')

        # Параметры для вставки значения
        table_index_target = 0  # Индекс таблицы в целевом документе
        row_index_target = 1     # Индекс строки в целевой таблице
        col_index_target = 0     # Индекс столбца в целевой таблице

        # Вставляем значение в целевой документ
        target_doc = 'sample.docx'
        new_file_name = f'{message.chat.id}_numbers.docx'
        act_number_to_insert = f'Номер уведомления о ВК / Акта ВК № {act_number[1]}'
        insert_value_to_table(target_doc, table_index_target, tab_rows_list[k], tab_cols_list[k], value_list[k], new_file_name)
        k = 0
        tab_rows_list = [2,4,5,6]
        tab_cols_list = [2,2,2,2]
        value_list = [act_number_to_insert, all_equipments_for_insert, all_kks_numbers_for_insert, all_safety_classes_for_insert, all_of_each_counts_for_insert]
        
        n = True
        while n == True:
            print(k, ' k')
            n = insert_value_to_table(new_file_name, table_index_target, tab_rows_list[k], tab_cols_list[k], value_list[k], new_file_name)
            k += 1
            if k == 4:
                n = False
            
        # insert_value_to_table(target_doc, table_index_target, row_index_target, col_index_target, value_to_insert, new_file_name)
        # insert_value_to_table(target_doc, table_index_target, 2, 2, all_equipments_for_insert, new_file_name)
        # insert_value_to_table(target_doc, table_index_target, 4, 2, all_kks_numbers_for_insert, new_file_name)
        # insert_value_to_table(target_doc, table_index_target, 5, 2, all_safety_classes_for_insert, new_file_name)
        # insert_value_to_table(target_doc, table_index_target, 6, 2, all_of_each_counts_for_insert, new_file_name)
        


        
        # Удаляем файл после обработки
        os.remove(message.document.file_name)

    else:
        bot.send_message(message.chat.id, 'Формат не Word .doc')

@bot.message_handler(commands=['show_final_document'])
def show(message):
    connection = db.create_connection("DEFECTS.sqlite", '3')
    #вывод таблицы defects
    getColumnNames=connection.execute("select * from defects limit 1")
    colName=[i[0] for i in getColumnNames.description]

    vivod=f'SELECT * FROM defects WHERE telegram_id={message.chat.id}'
    viv=db.execute_read_query(connection, vivod)
    connection.close()
    print(tabulate(viv, headers=colName, tablefmt='grid'))

@bot.message_handler(commands=['final'])
def final_document(message):

    # print(message.text)
    if message.text == '/final':

        # Запрос названия объекта, где происходит проверка инженером
        mess = bot.send_message(message.chat.id, 'Введите название объекта проверки')
        bot.register_next_step_handler(mess, final_document)
    else:

        # Здесь должен быть функционал для формирования таблицы Эксель

        connection = db.create_connection("DEFECTS.sqlite", '4')
        get_all_user_defects = f"SELECT defect FROM defects WHERE telegram_id={message.chat.id}"
        all_user_defects ,= db.execute_read_query(connection, get_all_user_defects, '5')


        all_user_defects_list = all_user_defects[0].split(sep='$$$') #sd


        list_length = len(all_user_defects_list)-1

        i = 0 # Итератор - по количеству строк
        while i < list_length:
            defect = all_user_defects_list[i]
            print(i+1, defect)
            i+=1

        # # Работа с Word
        

        # # Конец работы с Word


        clear_user_row = f"UPDATE defects SET defect=' ' WHERE telegram_id={message.chat.id}"
        db.execute_query(connection, clear_user_row, 'Очистка строки юзера')
        connection.close()

        with open(f'{message.chat.id}_numbers.docx', 'rb') as word:
            bot.send_document(message.chat.id, word) # отправка в чат конечного документа
        os.remove(f'{message.chat.id}_numbers.docx') # после отправки документ удаляется

        bot.send_message(message.chat.id, 'Для начала введите команду /start.')

bot.polling()
