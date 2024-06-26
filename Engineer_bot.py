import os
import telebot
import speech_recognition
from pydub import AudioSegment
import DB as db
from tabulate import tabulate
import openpyxl
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook


# Ниже нужно вставить токен, который дал BotFather при регистрации
# Пример: token = ''
token = '6416692197:AAG5sa9mOys1kwru9ISEvQwtyMuwqRloEVE'  # <<< Ваш токен

bot = telebot.TeleBot(token)


def oga2wav(filename):
    # Конвертация формата файлов
    new_filename = filename.replace('.oga', '.wav')
    audio = AudioSegment.from_file(filename)
    audio.export(new_filename, format='wav')
    return new_filename


def recognize_speech(oga_filename):
    # Перевод голоса в текст + удаление использованных файлов
    wav_filename = oga2wav(oga_filename)
    recognizer = speech_recognition.Recognizer()

    with speech_recognition.WavFile(wav_filename) as source:     
        wav_audio = recognizer.record(source)

    text = recognizer.recognize_google(wav_audio, language='ru')

    if os.path.exists(oga_filename):
        os.remove(oga_filename)

    if os.path.exists(wav_filename):
        os.remove(wav_filename)

    return text


def download_file(bot, file_id):
    # Скачивание файла, который прислал пользователь
    file_info = bot.get_file(file_id)
    downloaded_file = bot.download_file(file_info.file_path)
    filename = file_id + file_info.file_path
    filename = filename.replace('/', '_')
    with open(filename, 'wb') as f:
        f.write(downloaded_file)
    return filename


@bot.message_handler(commands=['start'])
def say_hi(message):
    # Функция, отправляющая "Привет" в ответ на команду /start
    bot.send_message(message.chat.id, f'Привет, {message.from_user.first_name}.\nЗапиши голосовое сообщение и получи расшифровку.\n\nЧтобы сформировать окончательный документ введи команду /final')

    # Проверка на наличие юзера в базе
    # работа с БД
    connection = db.create_connection("DEFECTS.sqlite", 'Подключение для проверки на сущ юзера')
    checkUser=f"SELECT COUNT(telegram_id) FROM defects WHERE telegram_id={message.from_user.id} GROUP BY telegram_id"
    ch=db.execute_read_query(connection,checkUser,'Поиск юзера в базе')
    if ch:
        print('Юзер найден')
    else:
        print('нету юзера')
        zapros=f"INSERT INTO defects ('telegram_id') VALUES ({message.chat.id})"
        db.execute_query(connection, zapros, 'создание записи в БД для юзера')
    connection.close()


# defects = dict() #Словарь для сбора замечаний, необходимо перевести в БД
# d = 0

@bot.message_handler(content_types=['voice'])
def transcript(message):
    bot.send_message(message.chat.id, 'Идет расшифровка...')
    # Функция, отправляющая текст в ответ на голосовое
    filename = download_file(bot, message.voice.file_id)
    text = recognize_speech(filename)
    bot.send_message(message.chat.id, text)
    
    #работа с БД
    connection = db.create_connection("DEFECTS.sqlite", '2')

    take_defect = f"SELECT * FROM defects WHERE telegram_id={message.chat.id}"
    text_defect ,= db.execute_read_query(connection, take_defect, '3')
    # print(text_defect[2])
    text_for_db = text +' $$$ ' + text_defect[2]

    input_defect = f"UPDATE defects SET defect='{text_for_db}' WHERE telegram_id={message.chat.id}"
    db.execute_query(connection, input_defect, 'внесение в БД замечание')

    connection.close()
    #конец работы с БД

    # global d
    # d += 1
    # defects[f'Замечание {d}'] = text

@bot.message_handler(commands=['show_final_document'])
def show(message):
    connection = db.create_connection("DEFECTS.sqlite", '3')
    #вывод таблицы defects
    getColumnNames=connection.execute("select * from defects limit 1")
    colName=[i[0] for i in getColumnNames.description]

    vivod=f'SELECT * FROM defects WHERE telegram_id={message.chat.id}'
    viv=db.execute_read_query(connection, vivod)
    connection.close()
    print(tabulate(viv, headers=colName, tablefmt='grid'))

@bot.message_handler(commands=['final'])
def final_document(message):

    # print(message.text)
    if message.text == '/final':
    
        # Запрос названия объекта, где происходит проверка инженером
        mess = bot.send_message(message.chat.id, 'Введите название объекта проверки')
        bot.register_next_step_handler(mess, final_document)
    else:

        # Здесь должен быть функционал для формирования таблицы Эксель

        connection = db.create_connection("DEFECTS.sqlite", '4')
        get_all_user_defects = f"SELECT defect FROM defects WHERE telegram_id={message.chat.id}"
        all_user_defects ,= db.execute_read_query(connection, get_all_user_defects, '5')
    

        all_user_defects_list = all_user_defects[0].split(sep='$$$') #sd


        list_length = len(all_user_defects_list)-1

        i = 0 # Итератор - по количеству строк
        while i < list_length:
            defect = all_user_defects_list[i]
            print(i+1, defect) 
            i+=1
        
        # # Работа с Word
        doc = docx.Document()
        # Название документа
        title_text = f'Акт № '
        title_1_text = f'комиссионной целевой проверки промышленной безопасности на объекте {message.text}'
        
        comission = f'''Члены комиссии:
Инженер ОТН - \t\t\t\t {message.chat.last_name} {message.chat.first_name}
Начальник {message.text} - \t\t\t\t
Механик {message.text} - 
'''
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_1 = doc.add_heading(title_1_text, 2)
        title_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(comission)
        # добавляем таблицу 3x3
        table = doc.add_table(rows = i+1, cols = 4)
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        # Заголовок таблицы

        title_cell_num = table.cell(0,0)
        title_cell_num.text = str('№ пп')
        title_cell_desc = table.cell(0,1)
        title_cell_desc.text = str('Описание дефекта')
        title_cell_otv = table.cell(0,2)
        title_cell_otv.text = str('Ответственный за устранение')
        title_cell_srok = table.cell(0,3)
        title_cell_srok.text = str('Срок устранения')

        # заполняем таблицу данными
        for row in range(i):
            row+=1
            for col in range(3):
                if col == 0:
                    # получаем ячейку таблицы
                    cell = table.cell(row, col)
                    # записываем в ячейку данные
                    cell.text = str(row) # порядковый номер
                elif col == 1:
                    # получаем ячейку таблицы
                    cell = table.cell(row, col)
                    # записываем в ячейку данные
                    index = i - row 
                    cell.text = str(all_user_defects_list[index]) # Текст замечания
                elif col == 2:
                    # получаем ячейку таблицы
                    cell = table.cell(row, col)
                    # записываем в ячейку данные
                    cell.text = str(f'Начальник {message.text}') # Ответственный - надо сделать, чтобы пользователь определял
                else:
                    cell = table.cell(row, col)
                    cell.text = str(' ') # место для внесения срока

        doc.save(f'{message.chat.id}_numbers.docx') # сохранения конечного документа

        # # Конец работы с Word


        clear_user_row = f"UPDATE defects SET defect=' ' WHERE telegram_id={message.chat.id}"
        db.execute_query(connection, clear_user_row, 'Очистка строки юзера')
        connection.close()

        with open(f'{message.chat.id}_numbers.docx', 'rb') as word:
            bot.send_document(message.chat.id, word) # отправка в чат конечного документа
        os.remove(f'{message.chat.id}_numbers.docx') # после отправки документ удаляется

        bot.send_message(message.chat.id, 'Для начала введите команду /start.')

bot.polling()